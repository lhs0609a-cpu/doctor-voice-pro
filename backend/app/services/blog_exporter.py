"""
Blog Exporter Service
ë¸”ë¡œê·¸ ë³µë¶™ìš© ë¬¸ì„œ ìƒì„± (DOCX, HTML)
"""

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Optional
import re
import io
import base64
from PIL import Image


class BlogExporter:
    """
    ë¸”ë¡œê·¸ ë³µë¶™ìš© ë¬¸ì„œ ìƒì„± ì„œë¹„ìŠ¤
    """

    def __init__(self):
        # ê°•ì¡° ìƒ‰ìƒ íŒ”ë ˆíŠ¸
        self.highlight_colors = {
            "primary": RGBColor(255, 107, 107),  # ë¹¨ê°•
            "secondary": RGBColor(74, 144, 226),  # íŒŒë‘
            "success": RGBColor(34, 197, 94),  # ì´ˆë¡
            "warning": RGBColor(251, 191, 36),  # ë…¸ë‘
            "info": RGBColor(168, 85, 247),  # ë³´ë¼
        }

    def export_to_docx(
        self,
        content: str,
        title: str = "",
        images: Optional[List[Dict]] = None,
        keywords: Optional[List[str]] = None,
        emphasis_phrases: Optional[List[str]] = None,
    ) -> io.BytesIO:
        """
        ì›Œë“œ ë¬¸ì„œ ìƒì„± (ë„¤ì´ë²„ ë¸”ë¡œê·¸ ë³µë¶™ìš©)

        Args:
            content: ë³¸ë¬¸ ë‚´ìš©
            title: ê¸€ ì œëª©
            images: ì´ë¯¸ì§€ ë¦¬ìŠ¤íŠ¸ [{"url": "...", "position": 100, "caption": "..."}]
            keywords: ê°•ì¡°í•  í‚¤ì›Œë“œ ë¦¬ìŠ¤íŠ¸
            emphasis_phrases: í•˜ì´ë¼ì´íŠ¸í•  ë¬¸êµ¬ ë¦¬ìŠ¤íŠ¸

        Returns:
            BytesIO ê°ì²´ (ì›Œë“œ íŒŒì¼)
        """
        doc = Document()

        # ê¸°ë³¸ ìŠ¤íƒ€ì¼ ì„¤ì • - ì˜ë¬¸ í°íŠ¸ ì´ë¦„ë§Œ ì‚¬ìš© (ì¸ì½”ë”© ë¬¸ì œ í•´ê²°)
        style = doc.styles['Normal']
        font = style.font
        # Arial Unicode MSëŠ” í•œê¸€ì„ í¬í•¨í•œ ë‹¤êµ­ì–´ ì§€ì›
        # ë˜ëŠ” Calibrië„ í•œê¸€ ì˜ í‘œì‹œë¨
        font.name = 'Arial'
        font.size = Pt(11)

        # 1. ì œëª©ì€ ì›Œë“œì— í¬í•¨í•˜ì§€ ì•ŠìŒ
        # ë„¤ì´ë²„ ë¸”ë¡œê·¸ëŠ” ì œëª© ì…ë ¥ì°½ì´ ë³„ë„ë¡œ ìˆìœ¼ë¯€ë¡œ,
        # ì›Œë“œ íŒŒì¼ì—ëŠ” ë³¸ë¬¸ë§Œ í¬í•¨í•˜ì—¬ ë³µì‚¬-ë¶™ì—¬ë„£ê¸° ì‹œ ê¹”ë”í•˜ê²Œ ì²˜ë¦¬
        # (ì œëª©ì€ ì‚¬ìš©ìê°€ ë„¤ì´ë²„ ë¸”ë¡œê·¸ ì œëª©ë€ì— ì§ì ‘ ì…ë ¥)

        # 2. ë³¸ë¬¸ ì²˜ë¦¬
        paragraphs = [p for p in content.split('\n\n') if p.strip()]
        images_list = images or []
        total_paras = len(paragraphs)

        # ì´ë¯¸ì§€ ì‚½ì… ìœ„ì¹˜ ê³„ì‚°: ë¬¸ë‹¨ì„ ê· ë“±í•˜ê²Œ ë‚˜ëˆ”
        if images_list and total_paras > 0:
            # ì˜ˆ: 10ê°œ ë¬¸ë‹¨, 4ê°œ ì´ë¯¸ì§€ -> 2, 4, 6, 8ë²ˆì§¸ ë¬¸ë‹¨ ë’¤ì— ì‚½ì…
            insert_interval = max(2, total_paras // (len(images_list) + 1))
            image_positions = []
            for i in range(len(images_list)):
                pos = (i + 1) * insert_interval
                if pos < total_paras:
                    image_positions.append(pos)
                else:
                    image_positions.append(total_paras)
        else:
            image_positions = []

        img_index = 0
        for para_index, para_text in enumerate(paragraphs):
            # ì¸ìš©êµ¬ ê°ì§€
            if para_text.strip().startswith('>'):
                self._add_quote(doc, para_text.replace('>', '').strip())
            # ì œëª© ê°ì§€ (##)
            elif para_text.strip().startswith('##'):
                heading_text = para_text.replace('#', '').strip()
                h2 = doc.add_heading(heading_text, level=2)
                h2.runs[0].font.color.rgb = RGBColor(52, 73, 94)
            # ë¦¬ìŠ¤íŠ¸ ê°ì§€
            elif para_text.strip().startswith(('- ', '* ', 'â€¢ ')):
                self._add_bullet_list(doc, para_text)
            # ì¼ë°˜ ë¬¸ë‹¨
            else:
                self._add_styled_paragraph(
                    doc, para_text, keywords, emphasis_phrases
                )

            # ì´ë¯¸ì§€ ì‚½ì… (ê³„ì‚°ëœ ìœ„ì¹˜ì—ì„œ)
            if img_index < len(image_positions) and para_index + 1 == image_positions[img_index]:
                if img_index < len(images_list):
                    self._add_image_to_doc(doc, images_list[img_index])
                img_index += 1

        # 3. ë‚¨ì€ ì´ë¯¸ì§€ ë§ˆì§€ë§‰ì— ì¶”ê°€
        while img_index < len(images_list):
            self._add_image_to_doc(doc, images_list[img_index])
            img_index += 1

        # íŒŒì¼ì„ ë©”ëª¨ë¦¬ì— ì €ì¥
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # íŒŒì¼ í¬ê¸° ê²€ì¦ (ë„¤ì´ë²„ ë¸”ë¡œê·¸ 5MB ì œí•œ)
        file_size = len(file_stream.getvalue())
        file_size_mb = file_size / (1024 * 1024)
        print(f"ğŸ“„ DOCX íŒŒì¼ í¬ê¸°: {file_size_mb:.2f}MB")

        if file_size_mb > 5.0:
            print(f"âš ï¸ ê²½ê³ : íŒŒì¼ í¬ê¸°ê°€ 5MBë¥¼ ì´ˆê³¼í–ˆìŠµë‹ˆë‹¤ ({file_size_mb:.2f}MB)")
            print("   ë„¤ì´ë²„ ë¸”ë¡œê·¸ ì—…ë¡œë“œ ì‹œ ë¬¸ì œê°€ ë°œìƒí•  ìˆ˜ ìˆìŠµë‹ˆë‹¤.")
        else:
            print(f"âœ… íŒŒì¼ í¬ê¸° OK: ë„¤ì´ë²„ ë¸”ë¡œê·¸ ì—…ë¡œë“œ ê°€ëŠ¥")

        return file_stream

    def _add_styled_paragraph(
        self,
        doc: Document,
        text: str,
        keywords: Optional[List[str]] = None,
        emphasis_phrases: Optional[List[str]] = None,
    ):
        """
        ìŠ¤íƒ€ì¼ì´ ì ìš©ëœ ë¬¸ë‹¨ ì¶”ê°€ (ë„¤ì´ë²„ ë¸”ë¡œê·¸ ìµœì í™”)
        - í…ìŠ¤íŠ¸ ìƒ‰ìƒ ëŒ€ì‹  í˜•ê´‘íœ ì‚¬ìš© (ë„¤ì´ë²„ê°€ ìƒ‰ìƒì„ ë¬´ì‹œí•˜ê¸° ë•Œë¬¸)
        - ë³¼ë“œ + í˜•ê´‘íœ ì¡°í•©ìœ¼ë¡œ ê°•ì¡°
        """
        p = doc.add_paragraph()

        # í‚¤ì›Œë“œì™€ ê°•ì¡° ë¬¸êµ¬ë¥¼ í•˜ë‚˜ì˜ ë¦¬ìŠ¤íŠ¸ë¡œ í†µí•©
        highlight_terms = []
        if keywords:
            highlight_terms.extend([(kw, 'keyword') for kw in keywords])
        if emphasis_phrases:
            highlight_terms.extend([(phrase, 'emphasis') for phrase in emphasis_phrases])

        # ê¸¸ì´ ìˆœìœ¼ë¡œ ì •ë ¬ (ê¸´ ê²ƒë¶€í„° - ì¤‘ë³µ ë§¤ì¹­ ë°©ì§€)
        highlight_terms.sort(key=lambda x: len(x[0]), reverse=True)

        if not highlight_terms:
            # ê°•ì¡°í•  ê²ƒì´ ì—†ìœ¼ë©´ ê·¸ëƒ¥ ì¶”ê°€
            p.add_run(text)
            return

        # í…ìŠ¤íŠ¸ë¥¼ ë¶„í• í•˜ì—¬ ìŠ¤íƒ€ì¼ ì ìš©
        remaining_text = text
        processed_text = ""

        while remaining_text:
            found_match = False

            for term, term_type in highlight_terms:
                if term.lower() in remaining_text.lower():
                    # ëŒ€ì†Œë¬¸ì ë¬´ì‹œí•˜ê³  ì°¾ê¸°
                    idx = remaining_text.lower().find(term.lower())

                    # ë§¤ì¹­ ì „ í…ìŠ¤íŠ¸
                    before = remaining_text[:idx]
                    if before:
                        run = p.add_run(before)

                    # ë§¤ì¹­ëœ í…ìŠ¤íŠ¸ (ì›ë³¸ ì¼€ì´ìŠ¤ ìœ ì§€)
                    matched = remaining_text[idx:idx + len(term)]
                    run = p.add_run(matched)

                    if term_type == 'keyword':
                        # í‚¤ì›Œë“œ: ë…¸ë€ í˜•ê´‘íœ + ë³¼ë“œ (ë„¤ì´ë²„ ë¸”ë¡œê·¸ ìµœì í™”)
                        run.bold = True
                        run.font.highlight_color = 7  # Yellow - ë„¤ì´ë²„ì—ì„œ ì˜ ë³´ì„
                    else:
                        # ê°•ì¡° ë¬¸êµ¬: ì£¼í™© í˜•ê´‘íœ + ë³¼ë“œ
                        run.font.highlight_color = 12  # Bright Green (ëˆˆì— ë„ê²Œ)
                        run.bold = True

                    # ë‚˜ë¨¸ì§€ í…ìŠ¤íŠ¸
                    remaining_text = remaining_text[idx + len(term):]
                    found_match = True
                    break

            if not found_match:
                # ë” ì´ìƒ ë§¤ì¹­ë˜ëŠ” ê²ƒì´ ì—†ìœ¼ë©´ ë‚˜ë¨¸ì§€ ì¶”ê°€
                p.add_run(remaining_text)
                break

    def _add_quote(self, doc: Document, text: str):
        """
        ì¸ìš©êµ¬ ì¶”ê°€
        """
        quote = doc.add_paragraph(text)
        quote.style = 'Intense Quote'

        # ì™¼ìª½ í…Œë‘ë¦¬ ì¶”ê°€ (ì¸ìš©êµ¬ ìŠ¤íƒ€ì¼)
        pPr = quote._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')  # ë‘ê»˜
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), '4A90E2')  # íŒŒë€ìƒ‰
        pBdr.append(left)
        pPr.append(pBdr)

        # ë°°ê²½ìƒ‰ (ì—°í•œ íšŒìƒ‰)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

    def _add_bullet_list(self, doc: Document, text: str):
        """
        ë¶ˆë¦¿ ë¦¬ìŠ¤íŠ¸ ì¶”ê°€
        """
        items = text.split('\n')
        for item in items:
            item_text = re.sub(r'^[\s\-\*\â€¢]+', '', item).strip()
            if item_text:
                doc.add_paragraph(item_text, style='List Bullet')

    def _convert_to_png(self, img_source, max_width: int = 1200, quality: int = 85) -> io.BytesIO:
        """
        ì´ë¯¸ì§€ë¥¼ ìµœì í™”í•˜ì—¬ PNG/JPEG í˜•ì‹ìœ¼ë¡œ ë³€í™˜ (ë„¤ì´ë²„ ë¸”ë¡œê·¸ 5MB ì œí•œ ëŒ€ì‘)

        ë„¤ì´ë²„ ë¸”ë¡œê·¸ëŠ” EMF, WebP ë“± ì¼ë¶€ í˜•ì‹ì„ ì§€ì›í•˜ì§€ ì•ŠìŒ
        ì´ë¯¸ì§€ ë¦¬ì‚¬ì´ì§• + ì••ì¶•ìœ¼ë¡œ íŒŒì¼ í¬ê¸°ë¥¼ 5MB ì´í•˜ë¡œ ìµœì í™”

        Args:
            img_source: ì´ë¯¸ì§€ ì†ŒìŠ¤ (íŒŒì¼ ê²½ë¡œ ë˜ëŠ” BytesIO)
            max_width: ìµœëŒ€ ë„ˆë¹„ (ê¸°ë³¸ê°’: 1200px - ë„¤ì´ë²„ ë¸”ë¡œê·¸ ìµœì )
            quality: JPEG í’ˆì§ˆ (1-95, ê¸°ë³¸ê°’: 85)

        Returns:
            ìµœì í™”ëœ ì´ë¯¸ì§€ BytesIO
        """
        try:
            # ì´ë¯¸ì§€ ì—´ê¸°
            if isinstance(img_source, str):
                # íŒŒì¼ ê²½ë¡œ
                img = Image.open(img_source)
            else:
                # BytesIO ë˜ëŠ” bytes
                img = Image.open(img_source)

            # 1. ì´ë¯¸ì§€ ë¦¬ì‚¬ì´ì§• (í° ì´ë¯¸ì§€ëŠ” ì¶•ì†Œ)
            original_width, original_height = img.size
            if original_width > max_width:
                # ë¹„ìœ¨ ìœ ì§€í•˜ë©° ë¦¬ì‚¬ì´ì§•
                ratio = max_width / original_width
                new_height = int(original_height * ratio)
                img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
                print(f"âœ… ì´ë¯¸ì§€ ë¦¬ì‚¬ì´ì§•: {original_width}x{original_height} â†’ {max_width}x{new_height}")

            # 2. RGBAë¥¼ RGBë¡œ ë³€í™˜ (íˆ¬ëª…ë„ ì²˜ë¦¬)
            if img.mode in ('RGBA', 'LA', 'P'):
                # í°ìƒ‰ ë°°ê²½ì— í•©ì„±
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')

            # 3. JPEGë¡œ ì €ì¥ (PNGë³´ë‹¤ í›¨ì”¬ ì‘ì€ íŒŒì¼ í¬ê¸°)
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=quality, optimize=True)
            output.seek(0)

            # íŒŒì¼ í¬ê¸° í™•ì¸
            file_size = len(output.getvalue())
            file_size_mb = file_size / (1024 * 1024)
            print(f"âœ… ì´ë¯¸ì§€ ìµœì í™” ì™„ë£Œ: {file_size_mb:.2f}MB (í’ˆì§ˆ: {quality})")

            # 4. ë§Œì•½ íŒŒì¼ì´ ë„ˆë¬´ í¬ë©´ (1MB ì´ìƒ) ë” ì••ì¶•
            if file_size > 1 * 1024 * 1024:  # 1MB
                print(f"âš ï¸ íŒŒì¼ì´ í¼ ({file_size_mb:.2f}MB), ì¶”ê°€ ì••ì¶• ì§„í–‰...")
                output = io.BytesIO()
                # í’ˆì§ˆì„ ë‚®ì¶°ì„œ ë‹¤ì‹œ ì €ì¥
                img.save(output, format='JPEG', quality=70, optimize=True)
                output.seek(0)
                file_size = len(output.getvalue())
                file_size_mb = file_size / (1024 * 1024)
                print(f"âœ… ì¶”ê°€ ì••ì¶• ì™„ë£Œ: {file_size_mb:.2f}MB (í’ˆì§ˆ: 70)")

            return output

        except Exception as e:
            print(f"âš ï¸ ì´ë¯¸ì§€ ë³€í™˜ ì‹¤íŒ¨: {e}")
            raise

    def _add_image_to_doc(self, doc: Document, img_data: Dict):
        """
        ì´ë¯¸ì§€ ì¶”ê°€ (JPEG í˜•ì‹ìœ¼ë¡œ ìµœì í™”í•˜ì—¬ ë„¤ì´ë²„ ë¸”ë¡œê·¸ 5MB ì œí•œ ëŒ€ì‘)

        Args:
            img_data: {"url": "...", "caption": "...", "width": 5}
        """
        try:
            img_stream = None

            # URLì´ base64ì¸ ê²½ìš°
            if img_data.get('url', '').startswith('data:image'):
                # base64 ë””ì½”ë”©
                img_base64 = img_data['url'].split(',')[1]
                img_bytes = base64.b64decode(img_base64)
                # PNGë¡œ ë³€í™˜
                img_stream = self._convert_to_png(io.BytesIO(img_bytes))

            elif img_data.get('path'):
                # ë¡œì»¬ íŒŒì¼ ê²½ë¡œ - PNGë¡œ ë³€í™˜
                img_stream = self._convert_to_png(img_data['path'])

            if img_stream:
                width = img_data.get('width', 5)
                doc.add_picture(img_stream, width=Inches(width))

            # ìº¡ì…˜ ì¶”ê°€
            if img_data.get('caption'):
                caption = doc.add_paragraph(img_data['caption'])
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.runs[0].font.size = Pt(9)
                caption.runs[0].font.color.rgb = RGBColor(128, 128, 128)

            # ì´ë¯¸ì§€ í›„ ë¹ˆ ì¤„
            doc.add_paragraph()

        except Exception as e:
            print(f"âŒ ì´ë¯¸ì§€ ì¶”ê°€ ì‹¤íŒ¨: {e}")
            # ì´ë¯¸ì§€ ì¶”ê°€ ì‹¤íŒ¨ ì‹œ ì„¤ëª…ë§Œ ì¶”ê°€
            p = doc.add_paragraph(f"[ì´ë¯¸ì§€: {img_data.get('caption', 'ì´ë¯¸ì§€')}]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def export_to_naver_html(
        self,
        content: str,
        title: str = "",
        keywords: Optional[List[str]] = None,
        emphasis_phrases: Optional[List[str]] = None,
    ) -> str:
        """
        ë„¤ì´ë²„ ë¸”ë¡œê·¸ ì „ìš© HTML ìƒì„±

        ë„¤ì´ë²„ ìŠ¤ë§ˆíŠ¸ì—ë””í„°ê°€ ì§€ì›í•˜ëŠ” íƒœê·¸ë§Œ ì‚¬ìš©:
        - <b>, <font color="">, <font size="">
        - <blockquote>, <ul>, <li>
        """
        html_parts = []

        # ì œëª©
        if title:
            html_parts.append(f'<h2><b>{title}</b></h2>')
            html_parts.append('<br>')

        # ë³¸ë¬¸ ì²˜ë¦¬
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            if not para.strip():
                continue

            # ì¸ìš©êµ¬
            if para.strip().startswith('>'):
                quote_text = para.replace('>', '').strip()
                html_parts.append(f'<blockquote>{quote_text}</blockquote>')

            # ì œëª©
            elif para.strip().startswith('##'):
                heading = para.replace('#', '').strip()
                html_parts.append(f'<h3><b>{heading}</b></h3>')

            # ë¦¬ìŠ¤íŠ¸
            elif para.strip().startswith(('- ', '* ')):
                items = para.split('\n')
                html_parts.append('<ul>')
                for item in items:
                    item_text = re.sub(r'^[\s\-\*]+', '', item).strip()
                    if item_text:
                        html_parts.append(f'<li>{item_text}</li>')
                html_parts.append('</ul>')

            # ì¼ë°˜ ë¬¸ë‹¨
            else:
                styled_text = self._apply_naver_styles(
                    para, keywords, emphasis_phrases
                )
                html_parts.append(f'<p>{styled_text}</p>')

        return '\n'.join(html_parts)

    def _apply_naver_styles(
        self,
        text: str,
        keywords: Optional[List[str]] = None,
        emphasis_phrases: Optional[List[str]] = None,
    ) -> str:
        """
        ë„¤ì´ë²„ í˜¸í™˜ ìŠ¤íƒ€ì¼ ì ìš©
        """
        result = text

        # í‚¤ì›Œë“œ ê°•ì¡° (ë¹¨ê°„ìƒ‰ + ë³¼ë“œ)
        if keywords:
            for kw in sorted(keywords, key=len, reverse=True):
                pattern = re.compile(re.escape(kw), re.IGNORECASE)
                result = pattern.sub(
                    f'<b><font color="#FF6B6B">{kw}</font></b>',
                    result
                )

        # ê°•ì¡° ë¬¸êµ¬ (ë…¸ë€ ë°°ê²½)
        if emphasis_phrases:
            for phrase in sorted(emphasis_phrases, key=len, reverse=True):
                pattern = re.compile(re.escape(phrase), re.IGNORECASE)
                result = pattern.sub(
                    f'<span style="background-color: yellow;"><b>{phrase}</b></span>',
                    result
                )

        return result


# ì‹±ê¸€í†¤ ì¸ìŠ¤í„´ìŠ¤
blog_exporter = BlogExporter()
