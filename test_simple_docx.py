"""
간단한 DOCX 생성 테스트
"""
from docx import Document
from docx.shared import Pt
import io

def test_simple_docx():
    """기본 DOCX 생성 테스트"""
    print("1. Document 생성...")
    doc = Document()

    print("2. 스타일 설정...")
    style = doc.styles['Normal']
    font = style.font

    # 테스트 1: 한글 폰트 이름
    try:
        font.name = '맑은 고딕'
        print("   ✅ 한글 폰트 이름 설정 성공: '맑은 고딕'")
    except Exception as e:
        print(f"   ❌ 한글 폰트 이름 설정 실패: {e}")
        # 영문 폰트로 변경
        font.name = 'Arial'
        print("   → 영문 폰트로 변경: 'Arial'")

    font.size = Pt(11)

    print("3. 한글 제목 추가...")
    try:
        heading = doc.add_heading("테스트 제목입니다", level=1)
        print("   ✅ 한글 제목 추가 성공")
    except Exception as e:
        print(f"   ❌ 한글 제목 추가 실패: {e}")
        return False

    print("4. 한글 본문 추가...")
    try:
        p = doc.add_paragraph("이것은 한글 본문입니다. 건강검진, 치료, 예방")
        print("   ✅ 한글 본문 추가 성공")
    except Exception as e:
        print(f"   ❌ 한글 본문 추가 실패: {e}")
        return False

    print("5. BytesIO로 저장...")
    try:
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_size = len(file_stream.getvalue())
        print(f"   ✅ BytesIO 저장 성공: {file_size} bytes")
    except Exception as e:
        print(f"   ❌ BytesIO 저장 실패: {e}")
        import traceback
        traceback.print_exc()
        return False

    print("6. 파일로 저장...")
    try:
        with open("test_simple.docx", "wb") as f:
            f.write(file_stream.getvalue())
        print("   ✅ 파일 저장 성공: test_simple.docx")
    except Exception as e:
        print(f"   ❌ 파일 저장 실패: {e}")
        return False

    return True

if __name__ == "__main__":
    import sys
    sys.path.insert(0, 'backend')

    print("="*60)
    print("간단한 DOCX 생성 테스트")
    print("="*60 + "\n")

    success = test_simple_docx()

    print("\n" + "="*60)
    if success:
        print("✅ 모든 테스트 성공!")
    else:
        print("❌ 테스트 실패!")
    print("="*60)
